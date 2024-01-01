import pandas as pd
from docx import Document


class WordWriter:
    def __init__(self, full_name: str):
        self.doc = Document()
        self.full_name = full_name

    def save(self):
        """
        保存
        """
        self.doc.save(self.full_name + ".docx")

    def heading(self, t, level):
        """
        标题
        """
        self.doc.add_heading(t, level)

    def table(self, df: pd.DataFrame):
        """
        添加表格
        """
        tbl = self.doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
        for i, col in enumerate(df.columns):
            tbl.rows[0].cells[i].text = col
        for idx, row in df.iterrows():
            for jdx, col in enumerate(row):
                tbl.rows[idx + 1].cells[jdx].text = str(col)

    def paragraph(self, paragraph: str):
        self.doc.add_paragraph(paragraph)


if __name__ == '__main__':
    df = pd.DataFrame({'姓名': ['张三', '李四'], '年龄': [23, 24]})
    writer = WordWriter('a')
    writer.paragraph('我是一个段落')
    writer.paragraph('我是一个段落')
    writer.table(df)
    writer.save()
